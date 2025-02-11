from docx import Document
import os
import pandas as pd
import numpy as np

def get_result_from_testsheet(file_path): 
    doc = Document(file_path)

    for i, paragraph in enumerate(doc.paragraphs):
        if "Test List" in paragraph.text: # Start of real 
            start = i
            break
    
    answers = []
    for i, paragraph in enumerate(doc.paragraphs[start+1:]):
        if "A:" in paragraph.text:
            num = paragraph.text.split(".")[0]
            answer = paragraph.text.split("A:")[-1].replace("_","").strip()
            answers.append(answer)
    return answers

def npc_get_result_from_testsheet():
    np.random.seed(0)
    npc_eme_idx = np.arange(141)
    np.random.shuffle(npc_eme_idx)
    npc_eme_sizes = [30,30,30,30,21]
    npc_eme_sublists = np.split(npc_eme_idx, np.cumsum(npc_eme_sizes))
    npc_eme_df = pd.read_csv("results/tag_result/0916_1007_tagged_task1_emergent_with_relations.csv")
    # file_paths = [f"results/human_interview_result/ester/npc_test_{i}-ester.docx" for i in (0,1,2,3,4)]
    file_paths = [f"results/human_interview_result/kyuna_npc/kp; npc_test_{i}.docx" for i in (0,1,2,3,4)]

    sublists = []

    for i, (npc_eme_sublist, file_path)in enumerate(zip(npc_eme_sublists, file_paths)):
        sublist_df = npc_eme_df.iloc[npc_eme_sublist,:]
        sublist_df["human_kyunapark_generated"] = get_result_from_testsheet(file_path)
        sublists.append(sublist_df)
    
    result_list = pd.concat(sublists, axis=0)
    result_list.to_csv('npc_kyunapark_all.csv')

if __name__ == "__main__":
    npc_get_result_from_testsheet()
